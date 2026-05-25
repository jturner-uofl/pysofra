"""DOCX rendering via ``python-docx``.

The renderer abstracts away every low-level python-docx call. Callers only
ever invoke :meth:`~pysofra.core.SofraTable.to_docx`; they never need to
build paragraphs, runs, or XML cells themselves.

Theme hints come from the active theme's ``docx`` dict (font, sizing,
header borders, outer borders, zebra striping).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from ..core.schema import HeaderRow, Row, SpanningHeader
from ..core.table import SofraTable
from ..themes.registry import resolve_theme


@dataclass
class DocxRenderer:
    """Write a SofraTable to a ``.docx`` file."""

    def write(self, table: SofraTable, path: Path) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except ImportError as e:  # pragma: no cover
            raise ImportError(
                "DOCX export requires python-docx. Install with `pip install python-docx`."
            ) from e

        theme = resolve_theme(table.theme_name)
        d = theme.docx
        font_name: str = d.get("font_name", "Calibri")
        font_size: int = int(d.get("font_size", 10))
        header_bold: bool = bool(d.get("header_bold", True))
        header_bottom_border: bool = bool(d.get("header_bottom_border", True))
        outer_border: bool = bool(d.get("outer_border", True))
        row_zebra: bool = bool(d.get("row_zebra", False))

        doc = Document()
        if table.caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cap.add_run(table.caption)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size + 1)

        plot = getattr(table, "inline_plot", None)
        if plot is not None and getattr(plot, "png_bytes", None) and \
                table.inline_svg_position == "above":
            _embed_png(doc, plot.png_bytes, plot.width_in)

        ncols = _ncols(table)
        n_header_rows = (1 if table.spanning_headers else 0) + len(table.headers)
        n_body_rows = len(table.rows)
        n_total_rows = n_header_rows + n_body_rows
        if n_total_rows == 0:
            n_total_rows = 1

        word_table = doc.add_table(rows=n_total_rows, cols=ncols)
        # Default is autofit-on; users can disable via .autofit(False).
        word_table.autofit = bool((table.metadata or {}).get("autofit", True))
        word_table.style = "Table Grid" if outer_border else "Normal Table"

        row_idx = 0
        if table.spanning_headers:
            _write_spanning_row(
                word_table, row_idx, table.spanning_headers, ncols,
                font_name=font_name, font_size=font_size, qn=qn,
                OxmlElement=OxmlElement, header_bold=header_bold,
            )
            row_idx += 1

        for hr in table.headers:
            _write_header_row(
                word_table, row_idx, hr, ncols,
                font_name=font_name, font_size=font_size,
                header_bold=header_bold,
                header_bottom_border=header_bottom_border and hr is table.headers[-1],
                qn=qn, OxmlElement=OxmlElement,
            )
            row_idx += 1

        for r_idx, body_row in enumerate(table.rows):
            zebra = row_zebra and (r_idx % 2 == 1)
            _write_body_row(
                word_table, row_idx, body_row, ncols,
                font_name=font_name, font_size=font_size,
                zebra=zebra, qn=qn, OxmlElement=OxmlElement,
            )
            row_idx += 1

        if plot is not None and getattr(plot, "png_bytes", None) and \
                table.inline_svg_position == "below":
            _embed_png(doc, plot.png_bytes, plot.width_in)

        if table.footnotes:
            for footnote in table.footnotes:
                para = doc.add_paragraph()
                run = para.add_run(footnote)
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(max(8, font_size - 1))

        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        # python-docx stamps every ZIP entry with the current wall-clock,
        # which breaks cross-process byte-determinism. Rewrite with fixed
        # entry mtimes so identical input always yields identical bytes.
        from ._zip_determinism import make_zip_deterministic
        make_zip_deterministic(path)
        return path

    # Renderer interface: render() returns the path string. Kept for
    # symmetry with the other renderers; most callers use to_docx().
    def render(self, table: SofraTable) -> str:  # pragma: no cover
        raise NotImplementedError("DocxRenderer writes to disk; use .write(table, path).")


# ----------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------


def _ncols(table: SofraTable) -> int:
    if table.headers:
        return len(table.headers[0].cells)
    if table.rows:
        return len(table.rows[0].cells)
    return 1


def _embed_png(doc: Any, png_bytes: bytes, width_in: float) -> None:
    """Insert a PNG image at the current insertion point."""
    import io

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(width_in))


def _set_cell_text(cell: Any, text: str, *, bold: bool, italic: bool,
                   font_name: str, font_size: int, align: str | None,
                   indent_em: float = 0.0,
                   parts: Any = None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = ""  # clear default empty paragraph contents
    para = cell.paragraphs[0]
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent_em > 0:
        from docx.shared import Pt as _Pt  # noqa: F401
        para.paragraph_format.left_indent = Pt(indent_em * 12)
    if parts:
        # One run per CellPart, each carrying its own formatting.
        for p in parts:
            run = para.add_run(p.text)
            run.bold = bool(p.bold) or bold
            run.italic = bool(p.italic) or italic
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if p.superscript:
                run.font.superscript = True
            if p.subscript:
                run.font.subscript = True
            if p.code:
                run.font.name = "Courier New"
            if p.color:
                try:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor.from_string(
                        p.color.lstrip("#")
                    )
                except Exception:  # pragma: no cover — bad colour string
                    pass
        return
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)


def _write_header_row(word_table: Any, idx: int, hr: HeaderRow, ncols: int,
                      *, font_name: str, font_size: int,
                      header_bold: bool, header_bottom_border: bool,
                      qn: Any, OxmlElement: Any) -> None:
    row = word_table.rows[idx]
    for j, cell in enumerate(hr.cells[:ncols]):
        wc = row.cells[j]
        _set_cell_text(
            wc,
            cell.text.replace("\n", "\n"),
            bold=header_bold or cell.bold,
            italic=False,
            font_name=font_name,
            font_size=font_size,
            align=cell.align,
        )
        if header_bottom_border:
            _set_cell_borders(wc, bottom="single", qn=qn, OxmlElement=OxmlElement, size=8)


def _write_body_row(word_table: Any, idx: int, r: Row, ncols: int,
                    *, font_name: str, font_size: int, zebra: bool,
                    qn: Any, OxmlElement: Any) -> None:
    row = word_table.rows[idx]
    for j, c in enumerate(r.cells[:ncols]):
        wc = row.cells[j]
        _set_cell_text(
            wc, c.text,
            bold=c.bold or r.is_group_header,
            italic=c.italic,
            font_name=font_name,
            font_size=font_size,
            align=c.align,
            indent_em=c.indent * 1.2,
            parts=c.parts,
        )
        if zebra:
            _set_cell_shading(wc, "F2F2F2", qn=qn, OxmlElement=OxmlElement)


def _write_spanning_row(word_table: Any, idx: int, spans: tuple[SpanningHeader, ...],
                        ncols: int, *, font_name: str, font_size: int,
                        qn: Any, OxmlElement: Any, header_bold: bool) -> None:
    row = word_table.rows[idx]
    # Set labels and merge.
    for span in spans:
        anchor = row.cells[span.start]
        for j in range(span.start + 1, span.end + 1):
            anchor = anchor.merge(row.cells[j])
        _set_cell_text(
            anchor, span.label,
            bold=header_bold, italic=False,
            font_name=font_name, font_size=font_size, align="center",
        )
        _set_cell_borders(anchor, bottom="single", qn=qn, OxmlElement=OxmlElement, size=4)


def _set_cell_borders(cell: Any, *, bottom: str | None = None, top: str | None = None,
                      size: int = 4, qn: Any, OxmlElement: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, style in (("bottom", bottom), ("top", top)):
        if style is None:
            continue
        existing = borders.find(qn(f"w:{edge}"))
        if existing is not None:
            borders.remove(existing)
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), "000000")
        borders.append(el)


def _set_cell_shading(cell: Any, color_hex: str, *, qn: Any, OxmlElement: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
